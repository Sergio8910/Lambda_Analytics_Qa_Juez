"""Parser de información previa: bytes + nombre de archivo -> ReferenceDataset.

Formatos soportados:
  - .xlsx        -> tabular (openpyxl); primera fila = encabezados
  - .csv / .tsv  -> tabular (csv stdlib; detecta delimitador)
  - .json        -> lista de objetos (records) o un objeto suelto
  - .txt / .md   -> texto plano
  - .docx        -> texto plano + tablas (python-docx)

No ejecuta nada ni hace red: solo lee el contenido en memoria.
"""
from __future__ import annotations

import csv
import io
import json
from pathlib import Path
from typing import Any, Dict, List

from .models import ReferenceDataset, ReferenceRecord


class ParseError(Exception):
    """El archivo no se pudo parsear (formato inválido o corrupto)."""


_TABULAR = {".xlsx", ".xls"}
_DELIMITED = {".csv", ".tsv"}
_TEXT = {".txt", ".md", ".text"}


def parse_reference_file(filename: str, content: bytes) -> ReferenceDataset:
    """Parsea un archivo de información previa al dataset estructurado neutral.

    Args:
        filename: nombre original (se usa la extensión para decidir el parser).
        content: bytes crudos del archivo.

    Raises:
        ParseError: si la extensión no se soporta o el contenido es inválido.
    """
    ext = Path(filename or "").suffix.lower()
    if ext in _TABULAR:
        return _parse_xlsx(filename, content)
    if ext in _DELIMITED:
        return _parse_delimited(filename, content, ext)
    if ext == ".json":
        return _parse_json(filename, content)
    if ext == ".docx":
        return _parse_docx(filename, content)
    if ext in _TEXT:
        return _parse_text(filename, content)
    raise ParseError(
        f"Extensión no soportada: '{ext or '(sin extensión)'}'. "
        f"Use: .xlsx, .csv, .tsv, .json, .txt o .docx."
    )


# ---------------------------------------------------------------------------
# Helpers comunes
# ---------------------------------------------------------------------------


def _rows_to_records(rows: List[List[Any]], notas: List[str]) -> tuple[List[str], List[ReferenceRecord]]:
    """Primera fila = encabezados; el resto = records."""
    if not rows:
        return [], []
    header = [str(c).strip() if c is not None else "" for c in rows[0]]
    # Rellena nombres de columna vacíos para no perder datos.
    columns: List[str] = []
    for i, name in enumerate(header):
        if name:
            columns.append(name)
        else:
            columns.append(f"col_{i + 1}")
            notas.append(f"Columna {i + 1} sin encabezado; se nombró 'col_{i + 1}'.")
    records: List[ReferenceRecord] = []
    for raw in rows[1:]:
        if all(c is None or str(c).strip() == "" for c in raw):
            continue  # fila totalmente vacía
        rec: ReferenceRecord = {}
        for i, col in enumerate(columns):
            rec[col] = raw[i] if i < len(raw) else None
        records.append(rec)
    return columns, records


def _decode(content: bytes) -> str:
    for enc in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            return content.decode(enc)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace")


# ---------------------------------------------------------------------------
# Parsers por formato
# ---------------------------------------------------------------------------


def _parse_xlsx(filename: str, content: bytes) -> ReferenceDataset:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:  # pragma: no cover
        raise ParseError("Falta 'openpyxl' para leer Excel. Instala: pip install openpyxl") from exc
    try:
        wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    except Exception as exc:
        raise ParseError(f"No se pudo abrir el Excel: {exc}") from exc

    notas: List[str] = []
    ws = wb.active
    if len(wb.sheetnames) > 1:
        notas.append(
            f"El libro tiene {len(wb.sheetnames)} hojas; se usó la activa "
            f"('{ws.title}'). Hojas: {', '.join(wb.sheetnames)}."
        )
    rows = [list(r) for r in ws.iter_rows(values_only=True)]
    wb.close()
    columns, records = _rows_to_records(rows, notas)
    return ReferenceDataset(
        source_name=filename, format="xlsx", columns=columns, records=records, notas=notas,
    )


def _parse_delimited(filename: str, content: bytes, ext: str) -> ReferenceDataset:
    text = _decode(content)
    notas: List[str] = []
    delimiter = "\t" if ext == ".tsv" else None
    if delimiter is None:
        try:
            dialect = csv.Sniffer().sniff(text[:4096], delimiters=",;\t|")
            delimiter = dialect.delimiter
        except csv.Error:
            delimiter = ","
            notas.append("No se detectó el delimitador; se asumió ','.")
    reader = csv.reader(io.StringIO(text), delimiter=delimiter)
    rows = [row for row in reader]
    columns, records = _rows_to_records(rows, notas)
    return ReferenceDataset(
        source_name=filename, format="csv", columns=columns, records=records, notas=notas,
    )


def _parse_json(filename: str, content: bytes) -> ReferenceDataset:
    try:
        data = json.loads(_decode(content))
    except json.JSONDecodeError as exc:
        raise ParseError(f"JSON inválido: {exc}") from exc

    notas: List[str] = []
    records: List[ReferenceRecord] = []
    if isinstance(data, list):
        records = [d if isinstance(d, dict) else {"value": d} for d in data]
    elif isinstance(data, dict):
        # Si hay una sola clave con una lista de objetos, úsala como records.
        list_keys = [k for k, v in data.items() if isinstance(v, list)]
        if len(list_keys) == 1 and all(isinstance(x, dict) for x in data[list_keys[0]]):
            records = data[list_keys[0]]
            notas.append(f"Se tomó la lista bajo la clave '{list_keys[0]}' como records.")
        else:
            records = [data]
    else:
        raise ParseError("El JSON debe ser un objeto o una lista de objetos.")

    columns: List[str] = []
    for rec in records:
        for k in rec.keys():
            if k not in columns:
                columns.append(k)
    return ReferenceDataset(
        source_name=filename, format="json", columns=columns, records=records, notas=notas,
    )


def _parse_docx(filename: str, content: bytes) -> ReferenceDataset:
    try:
        import docx  # python-docx
    except ImportError as exc:  # pragma: no cover
        raise ParseError("Falta 'python-docx' para leer Word. Instala: pip install python-docx") from exc
    try:
        doc = docx.Document(io.BytesIO(content))
    except Exception as exc:
        raise ParseError(f"No se pudo abrir el Word: {exc}") from exc

    notas: List[str] = []
    parrafos = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    text = "\n".join(parrafos)

    # Si hay tablas, la primera se interpreta como dataset tabular.
    columns: List[str] = []
    records: List[ReferenceRecord] = []
    if doc.tables:
        if len(doc.tables) > 1:
            notas.append(f"El documento tiene {len(doc.tables)} tablas; se usó la primera.")
        tabla = doc.tables[0]
        rows = [[cell.text.strip() for cell in row.cells] for row in tabla.rows]
        columns, records = _rows_to_records(rows, notas)
    return ReferenceDataset(
        source_name=filename, format="docx", columns=columns, records=records, text=text, notas=notas,
    )


def _parse_text(filename: str, content: bytes) -> ReferenceDataset:
    text = _decode(content)
    return ReferenceDataset(source_name=filename, format="txt", text=text)
